#!/usr/bin/env python3
"""Run the whole offline pipeline on the fixtures and assert the outcome.

This is the repo's smoke test: one command that answers "does the pipeline
still work". It needs no DEEPL_AUTH_KEY and no network — the three DeepL stages
(glossary push, translate, back-translate) are not exercised here and have to be
verified against the live API separately. `review_sources.py check` is out for
the same reason: its fetch path is network-only, so only the offline
`verify-manifest` mode is asserted here (docs/harness-backlog.md).

The assertions, all from CONTRACTS.md:
  - on the clean fixture, every check has its expected status. The one failure
    and the one warning are deliberate source quirks (see fixture_spec.md), so
    an all-pass run means the checks stopped working.
  - on the seeded-defect fixture, each of the six mutations is caught by the
    check the contract assigns to it.
  - the review packet has exactly the entries the fixture's segments call for,
    entry keys exactly {id, kind, section, text_it, text_en}, no timestamp and
    no hash, and NOT ONE forbidden field: the anti-anchoring rule is asserted
    against the project's own answer key, so every flag issue, check failure
    detail, edit reason, claim convention and undelivered DeepL variant must be
    absent from the packet bytes. That is the assertion the whole blind design
    rests on — a builder that leaks any of them fails this run.
  - review_packet.py prints the contract's state digest and leaves state/
    byte-identical, and two builds of the same state are byte-identical.
  - review_validate.py accepts a valid findings file, and every rejection reason
    in the contract fires on a vector of its own.
  - review_compare.py sorts canned findings and canned flags into the three
    expected buckets, in the contract's order, and leaves state/ untouched.
  - review_sources.py verify-manifest agrees with the style guide's Sources
    section, and disagreement is detected: a bumped edition, a dropped entry and
    a moved URL each fail.
  - the seeded review fixture carries all eight judgement defects, and each one
    arrives in the packet the reviewer actually reads.

Usage:
    smoke.py [--keep-going]

Run from the repo root. Exit 0 means clean.
"""

from __future__ import annotations

import argparse
import copy
import csv
import hashlib
import io
import json
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

SCRIPTS = Path(".claude/skills/patent-translate/scripts")
REFERENCES = Path(".claude/skills/patent-translate/references")
MANIFEST = REFERENCES / "review_sources.json"
STYLE_GUIDE = REFERENCES / "style-guide.md"
SEED_TABLE = SCRIPTS / "dev" / "review_seed_revisions.json"

FIXTURE = "projects/_fixture"
FIXTURE_BAD = "projects/_fixture_bad"
FIXTURE_REVIEW = "projects/_fixture_review"

# Clean fixture: exact status of every check. claim_support fails and
# numeral_term_consistency warns by design — both are seeded in the Italian.
EXPECTED_CLEAN = {
    "paragraph_parity": "pass",
    "numerals_per_segment": "pass",
    "numeral_term_consistency": "warn",
    "numbers_units": "pass",
    "claims_graph": "pass",
    "one_sentence_per_claim": "pass",
    "claim_support": "fail",
    "glossary_adherence": "pass",
    "antecedent_basis": "pass",
    "abstract_length": "pass",
    "residue": "pass",
    "structure": "pass",
    "decimal_comma_audit": "pass",
}

# Seeded-defect fixture: the six outcomes CONTRACTS.md requires. Other checks
# may also react (a term swap moves more than one check) and that is fine.
EXPECTED_BAD = {
    "numerals_per_segment": "fail",
    "numbers_units": "fail",
    "claim_support": "fail",
    "claims_graph": "fail",
    "one_sentence_per_claim": "fail",
    "antecedent_basis": "warn",
}

DELIVERABLES = (
    "filing_en.docx",
    "side_by_side.docx",
    "ESCALATIONS.md",
    "audit_numbers.csv",
    "terminology.csv",
    "bilingual.csv",
)

# --------------------------------------------------------------------------
# review packet

PACKET_FILES = ("pairs.json", "style-guide.md", "terminology.csv")
PACKET_ENTRY_KEYS = {"id", "kind", "section", "text_it", "text_en"}

# review_packet.PACKET_GLOSSARY_COLUMNS, restated here on purpose: smoke must
# fail if the builder's own idea of the packet columns drifts, so it cannot
# import the value it is checking.
PACKET_GLOSSARY_COLUMNS = ("term_it", "variants_it", "term_en", "variants_en",
                           "category", "in_claims")
# The glossary columns that carry the pipeline's doubts and must never reach the
# packet: `flag` is "empty when no doubt" by schema, `rationale` is the reasoning
# behind a choice, `status` says which rows were locked.
WITHHELD_GLOSSARY_COLUMNS = ("rationale", "status", "flag")

# The packet's golden structure: every segment of the fixture except the two
# untranslated section markers, in segments.json order. Ids, kinds and sections
# only — the texts are fixture prose and belong to fixture_spec.md, but the
# inventory is a contract: nothing sampled, nothing omitted, document order kept.
# projects/_fixture_review copies segments.json verbatim, so both packets are
# asserted against this one list.
EXPECTED_PACKET_ENTRIES = (
    ("T-001", "title", "title"),
    ("H-002", "heading", "abstract"),
    ("A-001", "abstract", "abstract"),
    ("H-003", "heading", "description"),
    ("H-004", "heading", "description"),
    ("D-0001", "description", "description"),
    ("H-005", "heading", "description"),
    ("D-0002", "description", "description"),
    ("H-006", "heading", "description"),
    ("D-0003", "description", "description"),
    ("D-0004", "description", "description"),
    ("H-007", "heading", "description"),
    ("D-0005", "description", "description"),
    ("D-0006", "description", "description"),
    ("D-0007", "description", "description"),
    ("D-0008", "description", "description"),
    ("D-0009", "description", "description"),
    ("D-0010", "description", "description"),
    ("C-01", "claim", "claims"),
    ("C-02", "claim", "claims"),
    ("C-03", "claim", "claims"),
    ("C-04", "claim", "claims"),
    ("C-05", "claim", "claims"),
    ("C-06", "claim", "claims"),
)
EXPECTED_PACKET_SECTIONS = {"title": 1, "abstract": 2, "description": 15, "claims": 6}
# H-001 (TITOLO) and H-008 (RIVENDICAZIONI) are section markers of the Italian
# source with no English anywhere in the pipeline: the packet omits them and says
# so on stdout.
EXPECTED_PACKET_OMITTED = "omitted 2 untranslated section marker(s): H-001, H-008"

# CONTRACTS.md: "No timestamp, no hash, no run id anywhere in the packet". Both
# patterns are absent from Italian and English patent prose, so a hit means a
# field was added to pairs.json, not that the document mentions a date.
TIMESTAMP_RE = re.compile(r"\d{4}-\d{2}-\d{2}T\d{2}:\d{2}:\d{2}")
HASH_RE = re.compile(r"\b[0-9a-f]{32,}\b")

# A forbidden string shorter than this is not searched for: a 4-character flag
# issue would match half the document and turn the anti-anchoring assertion into
# noise. Every answer-key string the fixtures actually carry is far longer.
MIN_ANSWER_KEY_LENGTH = 12

# --------------------------------------------------------------------------
# review_validate vectors

# The accept vector: one finding per interesting shape — optional keys present
# and absent, a DOC locus, a COVERAGE-GAP without a rule, a grammar finding with
# a proposal, an AMBIGUITY with its two readings, and an extra meta key (allowed).
VALID_FINDINGS = {
    "meta": {
        "project": "_fixture",
        "generated_at": "2026-07-30T09:12:00+00:00",
        "shard_plan": "smoke: one shard over the whole packet",
    },
    "findings": [
        {
            "id": "F-0001",
            "class": "TERM",
            "segment_id": "D-0002",
            "lens": "rules",
            "issue": "'sistemi' is rendered 'systems' in the first sentence and "
                     "left implicit in the second: one Italian term, two forms.",
            "options": [
                "render both occurrences as 'systems'",
                "render both occurrences as 'such systems'",
            ],
            "rule": "style-guide, Consistency is the deliverable",
            "proposed_en": None,
            "shard": "description-01-rules",
        },
        {
            "id": "F-0002",
            "class": "AMBIGUITY",
            "segment_id": "C-01",
            "lens": "rules",
            "issue": "the preamble transition admits an open and a closed reading",
            "options": [
                "open reading: 'comprendente' -> 'comprising'",
                "closed reading: 'consisting of', accepting the loss of scope",
            ],
            "rule": "style-guide, Transitional phrases",
        },
        {
            "id": "F-0003",
            "class": "COVERAGE-GAP",
            "segment_id": "DOC",
            "lens": "rules",
            "issue": "the guide does not regulate figure references ('la figura 1' "
                     "in D-0003 and D-0004): capitalised 'Figure 1' is practice",
            "options": ["add to Standard renderings: 'la figura N' -> 'Figure N'"],
        },
        {
            "id": "F-0004",
            "class": "CONVENTION",
            "segment_id": "D-0006",
            "lens": "grammar",
            "issue": "'with 100 µm mesh size' reads as a compound modifier with no "
                     "article",
            "options": ["'with a 100 µm mesh size'"],
            "proposed_en": "The filter assembly (3) comprises a metal mesh with a "
                           "100 µm mesh size and a removable support (3a).",
        },
    ],
}


def _v_unknown_key(doc, text_en):
    doc["findings"][0]["proposal_en"] = "a typo for proposed_en"
    return "unknown key 'proposal_en'"


def _v_missing_key(doc, text_en):
    del doc["findings"][0]["issue"]
    return "missing required key 'issue'"


def _v_bad_id(doc, text_en):
    doc["findings"][0]["id"] = "F-1"
    return "id 'F-1' does not match"


def _v_duplicate_id(doc, text_en):
    doc["findings"][1]["id"] = doc["findings"][0]["id"]
    return "id 'F-0001' is not unique"


def _v_mech_class(doc, text_en):
    doc["findings"][0]["class"] = "MECH"
    return "class 'MECH' is not one of"


def _v_bad_lens(doc, text_en):
    doc["findings"][0]["lens"] = "both"
    return "lens 'both' is not one of"


def _v_unknown_segment(doc, text_en):
    doc["findings"][0]["segment_id"] = "D-9999"
    return "segment_id 'D-9999' is not an entry id"


def _v_empty_segment(doc, text_en):
    doc["findings"][0]["segment_id"] = "   "
    return "segment_id must be a non-empty string"


def _v_empty_issue(doc, text_en):
    doc["findings"][0]["issue"] = "   "
    return "issue must be a non-empty string"


def _v_no_options(doc, text_en):
    doc["findings"][0]["options"] = []
    return "options must be a non-empty list"


def _v_option_placeholder(doc, text_en):
    doc["findings"][0]["options"] = ["To be reviewed with the firm.",
                                     "render both occurrences as 'systems'"]
    return "contains the placeholder phrase 'to be reviewed'"


def _v_option_bare(doc, text_en):
    doc["findings"][0]["options"] = ["Check.",
                                     "render both occurrences as 'systems'"]
    return "is the bare placeholder 'check'"


def _v_option_empty(doc, text_en):
    # Whitespace only: normalize() collapses it to "". A SHORT option is not a
    # violation — 'by' is the guide's own rendering of 'mediante'.
    doc["findings"][0]["options"] = ["   "]
    return "empty once normalized"


def _v_ambiguity_one_option(doc, text_en):
    doc["findings"][1]["options"] = doc["findings"][1]["options"][:1]
    return "class AMBIGUITY requires at least two options"


def _v_rules_without_rule(doc, text_en):
    del doc["findings"][0]["rule"]
    return "rule must be a non-empty string on a rules-lens finding"


def _v_proposed_noop(doc, text_en):
    doc["findings"][3]["proposed_en"] = text_en["D-0006"]
    return "proposed_en is identical to the delivered English of D-0006"


def _v_proposed_empty(doc, text_en):
    doc["findings"][3]["proposed_en"] = "   "
    return "proposed_en must be a non-empty string"


def _v_empty_shard(doc, text_en):
    doc["findings"][0]["shard"] = ""
    return "shard must be a non-empty string"


def _v_no_meta_project(doc, text_en):
    del doc["meta"]["project"]
    return "meta.project must be a non-empty string"


def _v_empty_generated_at(doc, text_en):
    doc["meta"]["generated_at"] = "   "
    return "meta.generated_at must be a non-empty string"


def _v_no_meta(doc, text_en):
    del doc["meta"]
    return "meta is missing or not an object"


def _v_findings_not_a_list(doc, text_en):
    doc["findings"] = {"F-0001": "…"}
    return "findings is missing or not a list"


# One vector per rejection rule of the review_validate.py contract. Each mutates
# a deep copy of the accept vector and returns the substring its own violation
# must contain, so the rule and its evidence stay in one place.
REJECT_VECTORS = (
    ("unknown finding key", _v_unknown_key),
    ("missing required key", _v_missing_key),
    ("id not F-dddd", _v_bad_id),
    ("duplicate id", _v_duplicate_id),
    ("MECH is not a reviewer class", _v_mech_class),
    ("lens outside rules/grammar", _v_bad_lens),
    ("segment_id not in the packet", _v_unknown_segment),
    ("segment_id blank", _v_empty_segment),
    ("issue blank", _v_empty_issue),
    ("options empty", _v_no_options),
    ("option is a placeholder phrase", _v_option_placeholder),
    ("option is a bare placeholder", _v_option_bare),
    ("option empty once normalized", _v_option_empty),
    ("AMBIGUITY with one option", _v_ambiguity_one_option),
    ("rules lens without a rule", _v_rules_without_rule),
    ("proposed_en is a no-op", _v_proposed_noop),
    ("proposed_en blank", _v_proposed_empty),
    ("shard blank", _v_empty_shard),
    ("meta.project absent", _v_no_meta_project),
    ("meta.generated_at blank", _v_empty_generated_at),
    ("meta absent", _v_no_meta),
    ("findings not a list", _v_findings_not_a_list),
)

# The three inputs that are not a mutated document: raw payload, None = no file.
RAW_REJECT_VECTORS = (
    ("top level is a JSON array", "[]\n", "expected a JSON object at the top level"),
    ("unparseable JSON", '{"meta": {"project": "_fixture"\n', "invalid JSON"),
    ("findings file absent", None, "file not found"),
)

# --------------------------------------------------------------------------
# review_compare canned inputs

# Five flags over four loci: two on one segment (so the cross product is
# exercised), one on a segment no finding touches, one document-wide (null
# segment_id), one on a segment the packet does not contain. One is already
# resolved and one is a CONVENTION: neither class nor status may filter.
CANNED_FLAGS = {
    "flags": [
        {"key": "aaaa11110001", "class": "TERM", "segment_id": "D-0002",
         "check": None, "text_it": None, "text_en": None,
         "issue": "canned flag A: 'sistemi' rendered inconsistently in D-0002",
         "options": ["render both occurrences as 'systems'"], "status": "open",
         "resolution": None, "created_by": "model",
         "created_at": "2026-07-30T09:00:00+00:00"},
        {"key": "bbbb11110002", "class": "CONVENTION", "segment_id": "D-0002",
         "check": None, "text_it": None, "text_en": None,
         "issue": "canned flag B: paragraph markers dropped in the English",
         "options": [], "status": "resolved",
         "resolution": "applied document-wide", "created_by": "model",
         "created_at": "2026-07-30T09:00:00+00:00"},
        {"key": "cccc11110003", "class": "CLAIM-DEFECT", "segment_id": "C-01",
         "check": None, "text_it": None, "text_en": None,
         "issue": "canned flag C: the preamble has no reference sign",
         "options": ["keep as filed and note it"], "status": "open",
         "resolution": None, "created_by": "model",
         "created_at": "2026-07-30T09:00:00+00:00"},
        {"key": "dddd11110004", "class": "AMBIGUITY", "segment_id": None,
         "check": None, "text_it": None, "text_en": None,
         "issue": "canned flag D: document-wide, the locus is a null segment_id",
         "options": ["reading one: …", "reading two: …"], "status": "open",
         "resolution": None, "created_by": "model",
         "created_at": "2026-07-30T09:00:00+00:00"},
        {"key": "eeee11110005", "class": "TERM", "segment_id": "H-009",
         "check": None, "text_it": None, "text_en": None,
         "issue": "canned flag E: on a segment the packet does not contain",
         "options": ["retranslate the heading"], "status": "open",
         "resolution": None, "created_by": "model",
         "created_at": "2026-07-30T09:00:00+00:00"},
    ]
}

CANNED_FINDINGS = {
    "meta": {"project": "canned", "generated_at": "2026-07-30T09:20:00+00:00"},
    "findings": [
        {"id": "F-0001", "class": "TERM", "segment_id": "D-0002", "lens": "rules",
         "issue": "canned finding 1: same segment as flags A and B",
         "options": ["render both occurrences as 'systems'"],
         "rule": "style-guide, Consistency is the deliverable"},
        {"id": "F-0002", "class": "CONVENTION", "segment_id": "D-0002",
         "lens": "grammar", "issue": "canned finding 2: same segment again",
         "options": ["'such systems have' -> 'such systems exhibit'"]},
        {"id": "F-0003", "class": "TERM", "segment_id": "D-0006", "lens": "rules",
         "issue": "canned finding 3: on a segment carrying no flag",
         "options": ["'mesh size' -> 'mesh aperture'"],
         "rule": "style-guide, Standard renderings"},
        {"id": "F-0004", "class": "COVERAGE-GAP", "segment_id": "DOC",
         "lens": "rules",
         "issue": "canned finding 4: document-wide, pairs with flag D",
         "options": ["add a rule on figure references"]},
        {"id": "F-0005", "class": "CLAIM-DEFECT", "segment_id": "C-06",
         "lens": "rules", "issue": "canned finding 5: on a segment carrying no flag",
         "options": ["restate the step in the gerund"],
         "rule": "style-guide, Claim structure"},
    ],
}

# (segment_id, finding_id, flag_key) in the contract's order: packet segment
# order, then finding id, then flag key, with the document-wide locus last.
EXPECTED_CANDIDATES = (
    ("D-0002", "F-0001", "aaaa11110001"),
    ("D-0002", "F-0001", "bbbb11110002"),
    ("D-0002", "F-0002", "aaaa11110001"),
    ("D-0002", "F-0002", "bbbb11110002"),
    ("DOC", "F-0004", "dddd11110004"),
)
EXPECTED_REVIEWER_ONLY = (("D-0006", "F-0003"), ("C-06", "F-0005"))
EXPECTED_PIPELINE_ONLY = (("C-01", "cccc11110003"), ("H-009", "eeee11110005"))
EXPECTED_COMPARE_COUNTS = {
    "findings": 5, "flags": 5, "candidates": 5,
    "reviewer_only": 2, "pipeline_only": 2,
}
# No flags is a normal input, not an error: every finding lands in reviewer_only,
# in packet order with the DOC finding last.
EXPECTED_NO_FLAGS_ORDER = ("F-0001", "F-0002", "F-0003", "F-0005", "F-0004")

# --------------------------------------------------------------------------
# review_sources vectors


def _m_bump_edition(doc):
    entry = _manifest_entry(doc, "pct-regulations")
    entry["style_guide_edition"] = "in force 1 January 2027"
    return "style_guide_edition 'in force 1 January 2027' does not occur"


def _m_drop_entry(doc):
    entry = _manifest_entry(doc, "cfr-37-1-75")
    doc["sources"] = [e for e in doc["sources"] if e is not entry]
    url = entry.get("style_guide_url") or entry["url"]
    return f"cites {url} in '## Sources' but no manifest entry claims it"


def _m_move_url(doc):
    entry = _manifest_entry(doc, "pct-articles")
    entry["style_guide_url"] = "https://example.invalid/moved.pdf"
    return "style_guide_url https://example.invalid/moved.pdf does not occur"


MANIFEST_REJECT_VECTORS = (
    ("edition bumped on one side only", _m_bump_edition),
    ("source dropped from the manifest", _m_drop_entry),
    ("cited URL moved on one side only", _m_move_url),
)

# --------------------------------------------------------------------------
# seeded review fixture

# The eight judgement failure modes projects/_fixture_review exists to carry,
# hardcoded here so a seed quietly dropped from the tracked seed table leaves
# this run red instead of leaving a failure mode unmeasured.
EXPECTED_SEED_MODES = (
    "means-to-mechanism",
    "reference-sign-relettered",
    "dropped-clause",
    "altered-defined-quantity",
    "gerunds-to-imperatives",
    "by-way-of",
    "closed-transition",
    "dropped-hedge",
)


class Failure(Exception):
    pass


# --------------------------------------------------------------------------
# running scripts


def run(script: str, *args: str) -> None:
    cmd = [sys.executable, str(SCRIPTS / script), *args]
    print(f"\n$ {' '.join(cmd[1:])}", flush=True)
    result = subprocess.run(cmd)
    if result.returncode != 0:
        raise Failure(f"{script} exited {result.returncode}")


def run_capture(script: str, *args: str, expect: int = 0) -> str:
    """Run a script, echo what it printed, and return its stdout."""
    cmd = [sys.executable, str(SCRIPTS / script), *args]
    print(f"\n$ {' '.join(cmd[1:])}", flush=True)
    result = subprocess.run(cmd, capture_output=True, text=True)
    if result.stdout:
        print(result.stdout, end="", flush=True)
    if result.stderr:
        print(result.stderr, end="", file=sys.stderr, flush=True)
    if result.returncode != expect:
        raise Failure(f"{script} exited {result.returncode}, expected {expect}")
    return result.stdout


def run_quiet(script: str, *args: str) -> tuple[int, str]:
    """(exit code, stdout + stderr), nothing echoed: the vector loops run the
    same script dozens of times and report one line per vector instead."""
    cmd = [sys.executable, str(SCRIPTS / script), *args]
    result = subprocess.run(cmd, capture_output=True, text=True)
    return result.returncode, result.stdout + result.stderr


# --------------------------------------------------------------------------
# pipeline assertions


def statuses(project: str) -> dict[str, str]:
    path = Path(project) / "state" / "checks.json"
    if not path.is_file():
        raise Failure(f"{path} not written")
    data = json.loads(path.read_text(encoding="utf-8"))
    return {r["check"]: r["status"] for r in data["results"]}


def details(project: str, check: str) -> list[str]:
    path = Path(project) / "state" / "checks.json"
    data = json.loads(path.read_text(encoding="utf-8"))
    for result in data["results"]:
        if result["check"] == check:
            return [f["detail"] for f in result["failures"]]
    return []


def assert_statuses(project: str, expected: dict[str, str], label: str) -> None:
    actual = statuses(project)
    problems = []
    for check, want in expected.items():
        got = actual.get(check)
        if got != want:
            problems.append(f"  {check}: expected {want}, got {got}")
    if problems:
        raise Failure(f"{label} check statuses differ:\n" + "\n".join(problems))
    print(f"OK: {label} — {len(expected)} check status(es) as expected")


def assert_deliverables() -> None:
    missing = [
        name for name in DELIVERABLES
        if not (Path(FIXTURE) / "out" / name).is_file()
    ]
    if missing:
        raise Failure(f"missing deliverables in {FIXTURE}/out: {', '.join(missing)}")

    # Existence alone is not an assertion: a stale file left by an earlier run
    # passes it. Read bilingual.csv back and check it against the state it is
    # supposed to have been built from.
    import csv as _csv

    segments = json.loads(
        (Path(FIXTURE) / "state" / "segments.json").read_text(encoding="utf-8")
    )["segments"]
    bilingual = Path(FIXTURE) / "out" / "bilingual.csv"
    with open(bilingual, encoding="utf-8-sig", newline="") as fh:
        rows = list(_csv.reader(fh))
    if rows[0] != ["id", "kind", "section", "italiano", "english"]:
        raise Failure(f"bilingual.csv header is {rows[0]}")
    if len(rows) - 1 != len(segments):
        raise Failure(
            f"bilingual.csv has {len(rows) - 1} rows, expected {len(segments)}"
        )
    if [r[0] for r in rows[1:]] != [s["id"] for s in segments]:
        raise Failure("bilingual.csv rows are not the segments in document order")
    # Section markers (TITOLO, RIVENDICAZIONI) carry no English by design and are
    # the only rows allowed to be empty; every translated segment must have text.
    empty = [r[0] for r in rows[1:] if r[1] != "heading" and not r[4].strip()]
    if empty:
        raise Failure("bilingual.csv has empty English on: " + ", ".join(empty))
    if bilingual.read_bytes()[:3] != b"\xef\xbb\xbf":
        raise Failure("bilingual.csv has no UTF-8 BOM; Excel will mangle the accents")

    print(f"OK: {len(DELIVERABLES)} deliverables written to {FIXTURE}/out, "
          f"bilingual.csv = {len(segments)} segments in document order")


NOTES_FOR_REVIEWER = "notes-for-human-reviewer.md"

# Filed on C-01 by the clean-fixture run so the exclusion below is testable.
FIXTURE_CONVENTION_ISSUE = (
    "Transitional phrase: 'comprendente' rendered 'comprising' throughout"
)


def assert_conventions_stay_internal() -> None:
    """CONVENTION reaches no deliverable, and the notes file is the only thing a
    human editor is told about drafting conventions.

    The fixture's flags.json carries a CONVENTION flag, so the absence asserted
    here is a real exclusion and not an empty set.
    """
    from docx import Document  # local: python-docx is only needed for this check

    flags = json.loads(
        (Path(FIXTURE) / "state" / "flags.json").read_text(encoding="utf-8"))
    conventions = [f for f in flags["flags"] if f["class"] == "CONVENTION"]
    on_a_segment = [f for f in conventions if f["segment_id"]]
    if not on_a_segment:
        raise Failure(
            "the fixture has no CONVENTION flag ON A SEGMENT, so this assertion "
            "would pass vacuously: a document-wide flag never reaches a "
            "review-table row whether CONVENTION is excluded or not")

    sbs = Document(str(Path(FIXTURE) / "out" / "side_by_side.docx"))
    body = "\n".join(p.text for p in sbs.paragraphs)
    table = "\n".join(c.text for r in sbs.tables[0].rows for c in r.cells)
    for flag in conventions:
        for where, blob in (("review table", table), ("appendices", body)):
            if flag["issue"] and flag["issue"] in blob:
                raise Failure(
                    f"CONVENTION flag {flag['key']} appears in the {where} of "
                    "side_by_side.docx; it is an internal record")
    if "Applied conventions" in body:
        raise Failure("side_by_side.docx still has an 'Applied conventions' appendix")
    escalations = (Path(FIXTURE) / "out" / "ESCALATIONS.md").read_text(encoding="utf-8")
    for flag in conventions:
        if flag["issue"] and flag["issue"] in escalations:
            raise Failure(f"CONVENTION flag {flag['key']} reached ESCALATIONS.md")

    # The notes file: copied when the project has one, and a stale out/ copy
    # removed when it does not. The fixture ships without one.
    source = Path(FIXTURE) / NOTES_FOR_REVIEWER
    target = Path(FIXTURE) / "out" / NOTES_FOR_REVIEWER
    if source.exists():
        raise Failure(f"the fixture must not ship a {NOTES_FOR_REVIEWER}")
    if target.exists():
        raise Failure(f"{target} exists although the project has no {NOTES_FOR_REVIEWER}")
    source.write_text("# Notes\n\n| Italian | English |\n| --- | --- |\n"
                      "| comprendente | comprising |\n", encoding="utf-8")
    try:
        code, output = run_quiet("assemble.py", "run", "--project", FIXTURE)
        if code != 0:
            raise Failure(f"assemble.py exited {code} while re-running:\n{output}")
        if target.read_text(encoding="utf-8") != source.read_text(encoding="utf-8"):
            raise Failure(f"{NOTES_FOR_REVIEWER} was not copied verbatim into out/")
        source.unlink()
        code, output = run_quiet("assemble.py", "run", "--project", FIXTURE)
        if code != 0:
            raise Failure(f"assemble.py exited {code} while re-running:\n{output}")
        if target.exists():
            raise Failure(
                f"out/{NOTES_FOR_REVIEWER} survived the project file being removed; "
                "a stale copy is worse than none")
    finally:
        source.unlink(missing_ok=True)
    print(f"OK: {len(conventions)} CONVENTION flag(s) reach no deliverable, no "
          f"'Applied conventions' appendix, and {NOTES_FOR_REVIEWER} is copied "
          "when present and removed when not")


# --------------------------------------------------------------------------
# review packet: shape, determinism, blindness


def state_digest(state_dir: Path) -> str:
    """The state digest of CONTRACTS.md, re-implemented here on purpose: the
    printed value has to be checked against the contract, not against the code
    that produced it."""
    digest = hashlib.sha256()
    files = sorted((p for p in state_dir.iterdir() if p.is_file()),
                   key=lambda p: p.name)
    for path in files:
        digest.update(path.name.encode("utf-8"))
        digest.update(b"\0")
        digest.update(path.read_bytes())
        digest.update(b"\0")
    return digest.hexdigest()


def packet_dir(project: str) -> Path:
    return Path(project) / "review" / "packet"


def packet_bytes(project: str) -> dict[str, bytes]:
    """Every file in the packet directory, so an extra one is visible too."""
    return {
        path.name: path.read_bytes()
        for path in sorted(packet_dir(project).iterdir())
        if path.is_file()
    }


def packet_text(project: str) -> str:
    """The whole packet as one string: what a blind pass can read."""
    return "\n".join(
        payload.decode("utf-8") for _, payload in sorted(packet_bytes(project).items())
    )


def assert_packet_build(project: str, label: str) -> None:
    """Build the packet and assert the three files, the omission line, and that
    state/ came out byte-identical with the digest the builder printed."""
    state_dir = Path(project) / "state"
    before = state_digest(state_dir)
    output = run_capture("review_packet.py", "build", "--project", project)

    lines = output.strip().splitlines()
    if not lines or lines[-1] != f"state_sha256 {before}":
        raise Failure(
            f"{label}: last line must be the contract's state digest "
            f"'state_sha256 {before}', got {lines[-1] if lines else '(no output)'!r}"
        )
    after = state_digest(state_dir)
    if after != before:
        raise Failure(
            f"{label}: building the packet changed {state_dir} "
            f"({before} -> {after}); the blind phase must not write to state/"
        )
    if EXPECTED_PACKET_OMITTED not in output:
        raise Failure(
            f"{label}: expected the builder to report '{EXPECTED_PACKET_OMITTED}'"
        )

    present = packet_bytes(project)
    if sorted(present) != sorted(PACKET_FILES):
        raise Failure(
            f"{label}: packet must hold exactly {', '.join(PACKET_FILES)}, "
            f"found {', '.join(sorted(present)) or '(nothing)'}"
        )
    if present["style-guide.md"] != STYLE_GUIDE.read_bytes():
        raise Failure(f"{label}: packet style-guide.md is not a byte copy of its source")
    for violation in glossary_projection_violations(project, present["terminology.csv"]):
        raise Failure(f"{label}: {violation}")
    print(f"OK: {label} — 3 packet files, style guide byte-identical, glossary "
          f"projected to {len(PACKET_GLOSSARY_COLUMNS)} columns, "
          f"state/ unchanged ({before[:12]}…)")


def glossary_projection_violations(project: str, packet_csv: bytes) -> list[str]:
    """The packet glossary must be the delivered one MINUS the doubt columns.

    Not a byte copy: `rationale`, `status` and `flag` are withheld because the
    glossary states both what was decided and whether we doubted it, and `flag`
    is "empty when no doubt" by its own schema — on cafe124 fifteen rows of it
    name the open TERM escalations outright. Two ways to break this, and both are
    checked: dropping a column the reviewer needs (it then cannot judge PCT Rule
    10.2 consistency at all), and carrying one it must not see. The withheld
    CONTENT is additionally covered by answer_key(), which scans the whole packet.
    """
    out: list[str] = []
    delivered = Path(project) / "out" / "terminology.csv"
    with delivered.open(newline="", encoding="utf-8") as handle:
        want_rows = list(csv.DictReader(handle))
    got_rows = list(csv.DictReader(io.StringIO(packet_csv.decode("utf-8"), newline="")))
    header = list(got_rows[0]) if got_rows else []
    if header != list(PACKET_GLOSSARY_COLUMNS):
        out.append(f"packet glossary columns are {header}, "
                   f"want {list(PACKET_GLOSSARY_COLUMNS)}")
        return out
    if len(got_rows) != len(want_rows):
        out.append(f"packet glossary has {len(got_rows)} rows, "
                   f"delivered has {len(want_rows)}")
        return out
    for index, (got, want) in enumerate(zip(got_rows, want_rows)):
        for column in PACKET_GLOSSARY_COLUMNS:
            if got[column] != want[column]:
                out.append(f"packet glossary row {index} column {column}: "
                           f"{got[column]!r} != delivered {want[column]!r}")
    return out


def packet_shape_violations(doc) -> list[str]:
    """Key sets of pairs.json, exactly. A missing key and a sixth key are both
    violations: the entry schema is closed, which is what keeps the packet from
    growing a field that anchors the reviewer."""
    out: list[str] = []
    if set(doc) != {"meta", "entries"}:
        out.append(f"top-level keys are {sorted(doc)}, want ['entries', 'meta']")
    meta = doc.get("meta")
    if not isinstance(meta, dict):
        out.append("meta is missing or not an object")
    else:
        if set(meta) != {"project", "counts"}:
            out.append(f"meta keys are {sorted(meta)}, want ['counts', 'project']")
        counts = meta.get("counts")
        if not isinstance(counts, dict):
            out.append("meta.counts is missing or not an object")
        else:
            if set(counts) != {"entries", "by_section"}:
                out.append(f"meta.counts keys are {sorted(counts)}, "
                           "want ['by_section', 'entries']")
            by_section = counts.get("by_section")
            if not isinstance(by_section, dict):
                out.append("meta.counts.by_section is missing or not an object")
            elif by_section != EXPECTED_PACKET_SECTIONS:
                out.append(f"meta.counts.by_section is {by_section}, "
                           f"want {EXPECTED_PACKET_SECTIONS}")
    entries = doc.get("entries")
    if not isinstance(entries, list):
        out.append("entries is missing or not a list")
        return out
    for index, entry in enumerate(entries):
        if not isinstance(entry, dict):
            out.append(f"entries[{index}] is not an object")
        elif set(entry) != PACKET_ENTRY_KEYS:
            out.append(
                f"entries[{index}] ({entry.get('id')}) keys are {sorted(entry)}, "
                f"want {sorted(PACKET_ENTRY_KEYS)}"
            )
    return out


def assert_packet_shape(project: str, label: str) -> None:
    payload = (packet_dir(project) / "pairs.json").read_text(encoding="utf-8")
    doc = json.loads(payload)
    problems = packet_shape_violations(doc)

    inventory = tuple(
        (e.get("id"), e.get("kind"), e.get("section")) for e in doc.get("entries", [])
    )
    if inventory != EXPECTED_PACKET_ENTRIES:
        extra = [e for e in inventory if e not in EXPECTED_PACKET_ENTRIES]
        missing = [e for e in EXPECTED_PACKET_ENTRIES if e not in inventory]
        problems.append(
            f"entry inventory differs: {len(inventory)} entries, "
            f"missing {missing or 'none'}, unexpected {extra or 'none'}"
        )
    counts = doc.get("meta", {}).get("counts", {})
    if counts.get("entries") != len(EXPECTED_PACKET_ENTRIES):
        problems.append(f"meta.counts.entries is {counts.get('entries')}, "
                        f"want {len(EXPECTED_PACKET_ENTRIES)}")
    for name, pattern in (("timestamp", TIMESTAMP_RE), ("hash", HASH_RE)):
        hit = pattern.search(payload)
        if hit is not None:
            problems.append(
                f"pairs.json carries a {name} ({hit.group(0)!r}): the packet must be "
                "byte-identical for identical state"
            )
    if problems:
        raise Failure(f"{label} packet shape:\n  " + "\n  ".join(problems))
    print(f"OK: {label} — {len(inventory)} entries in document order, "
          "entry keys exactly id/kind/section/text_it/text_en, no timestamp, no hash")


def answer_key(project: str) -> list[tuple[str, str]]:
    """(label, string) for every text in this project's state that would anchor a
    reviewer. None of them may appear anywhere in the packet.

    Only fields that are the pipeline's own conclusions are collected. A flag's
    `text_it` and `text_en` are the document itself and are in the packet by
    design, and terminology_report.json is a list of glossary terms that the
    packet's terminology.csv legitimately contains: including either would make
    this assertion fire on correct behaviour.

    A flag's `options` are excluded for the same reason, and it is not a
    weakening: an option IS a candidate English rendering, so by construction it
    may occur in the delivered English. On cafe124 five of them do — 'suitable
    for', 'provided with', 'equipped with', 'configured to', 'positionable' —
    and counting those as leaks would have made the assertion fire on a correct
    packet. What identifies a leak is the pipeline's PROSE ABOUT a doubt (the
    flag's issue and resolution, a check's failure detail, an edit's reason), not
    the words it proposes; and a packet that leaked flags.json would still be
    caught, by those issues.
    """
    root = Path(project)
    out: list[tuple[str, str]] = []

    def add(label: str, value) -> None:
        if isinstance(value, str) and len(value) >= MIN_ANSWER_KEY_LENGTH:
            out.append((label, value))

    flags_path = root / "state" / "flags.json"
    if flags_path.is_file():
        for flag in json.loads(flags_path.read_text(encoding="utf-8"))["flags"]:
            where = f"flags.json {flag.get('key')}"
            add(f"{where} key", flag.get("key"))
            add(f"{where} issue", flag.get("issue"))
            add(f"{where} resolution", flag.get("resolution"))

    checks_path = root / "state" / "checks.json"
    if checks_path.is_file():
        for result in json.loads(checks_path.read_text(encoding="utf-8"))["results"]:
            for failure in result["failures"]:
                add(f"checks.json {result['check']} detail", failure.get("detail"))

    translations_path = root / "state" / "translations.json"
    if translations_path.is_file():
        document = json.loads(translations_path.read_text(encoding="utf-8"))
        for pair in document["pairs"]:
            where = f"translations.json {pair['id']}"
            add(f"{where} deepl_para_sha256", pair.get("deepl_para_sha256"))
            for edit in pair.get("edits") or []:
                add(f"{where} edit reason", edit.get("reason"))
            # The DeepL variant is only an answer key where a final replaced it:
            # it then marks exactly which segments the pipeline doubted, and its
            # text is the reading the pipeline rejected.
            if pair.get("text_en_final") is not None:
                add(f"{where} undelivered text_en_deepl", pair.get("text_en_deepl"))

    # The glossary is legitimate packet content, but two of its columns are the
    # pipeline's doubts in prose and the packet withholds them. This is the leak
    # the file-name-based FORBIDDEN list could not see: terminology.csv was on
    # the allowed list, so nobody checked what was inside it.
    delivered_glossary = root / "out" / "terminology.csv"
    if delivered_glossary.is_file():
        with delivered_glossary.open(newline="", encoding="utf-8") as handle:
            for row in csv.DictReader(handle):
                where = f"out/terminology.csv {row.get('term_it')}"
                for column in WITHHELD_GLOSSARY_COLUMNS:
                    add(f"{where} {column}", row.get(column))

    claims_path = root / "state" / "claims_en.json"
    if claims_path.is_file():
        for claim in json.loads(claims_path.read_text(encoding="utf-8"))["claims"]:
            where = f"claims_en.json {claim['id']}"
            for convention in claim.get("conventions") or []:
                add(f"{where} convention", convention)
            add(f"{where} authored_at", claim.get("authored_at"))
            for edit in claim.get("edits") or []:
                add(f"{where} edit reason", edit.get("reason"))
                add(f"{where} superseded text", edit.get("before"))
    return out


def assert_packet_blind(project: str, label: str) -> None:
    """The anti-anchoring rule, asserted against this project's own answer key."""
    keys = answer_key(project)
    if not keys:
        raise Failure(
            f"{label}: no answer-key text found in {project}/state — the "
            "anti-anchoring assertion would pass vacuously"
        )
    text = packet_text(project)
    leaks = [f"{where}: {value[:70]!r}" for where, value in keys if value in text]
    if leaks:
        raise Failure(
            f"{label}: the packet leaks {len(leaks)} of {len(keys)} answer-key "
            "text(s), which is the one thing it may never do:\n  "
            + "\n  ".join(leaks)
        )
    print(f"OK: {label} — none of {len(keys)} answer-key texts "
          "(flag issues, check details, edit reasons, conventions, undelivered "
          "DeepL variants) appears in the packet")


def assert_packet_deterministic(project: str, label: str) -> None:
    before = packet_bytes(project)
    run_capture("review_packet.py", "build", "--project", project)
    after = packet_bytes(project)
    differing = sorted(
        name for name in set(before) | set(after)
        if before.get(name) != after.get(name)
    )
    if differing:
        raise Failure(
            f"{label}: rebuilding the same state changed {', '.join(differing)}; "
            "the packet must be byte-identical for identical state"
        )
    print(f"OK: {label} — two builds of the same state are byte-identical")


def assert_packet_refuses_stray_file(project: str, label: str) -> None:
    """A file the builder does not own must stop the build.

    A blind pass reads the packet DIRECTORY, so a leftover in it is part of the
    packet whatever it is. The builder never deletes project data, so it has to
    refuse; this vector plants the leftover, asserts the refusal names it, and
    removes it again so the packet is left as the earlier assertions found it.
    """
    stray = packet_dir(project) / "LEFTOVER.md"
    stray.write_text("an earlier run's notes, or an answer key\n", encoding="utf-8")
    try:
        code, output = run_quiet("review_packet.py", "build", "--project", project)
    finally:
        stray.unlink()
    if code != 1 or stray.name not in output:
        raise Failure(
            f"{label}: a stray {stray.name} in the packet directory must make the "
            f"build exit 1 naming the file; got exit {code} and:\n{output}"
        )
    if packet_bytes(project).keys() != set(PACKET_FILES):
        raise Failure(
            f"{label}: the refused build changed the packet directory contents"
        )
    print(f"OK: {label} — a stray file in the packet directory refuses the build "
          "(exit 1, file named), and the packet is untouched")


# --------------------------------------------------------------------------
# review_validate


def packet_text_en(project: str) -> dict[str, str]:
    doc = json.loads((packet_dir(project) / "pairs.json").read_text(encoding="utf-8"))
    return {entry["id"]: entry["text_en"] for entry in doc["entries"]}


def assert_validator(project: str, scratch: Path) -> None:
    """One accept vector, then one reject vector per rejection rule, each
    asserted on its exit code and on its own violation actually firing."""
    scratch.mkdir(parents=True, exist_ok=True)
    text_en = packet_text_en(project)

    accept = scratch / "accept.json"
    accept.write_text(json.dumps(VALID_FINDINGS, ensure_ascii=False, indent=2),
                      encoding="utf-8")
    code, output = run_quiet("review_validate.py", "--project", project,
                             "--input", str(accept))
    expected_ok = f"OK: {len(VALID_FINDINGS['findings'])} finding(s) valid"
    if code != 0 or expected_ok not in output:
        raise Failure(
            f"the accept vector was rejected (exit {code}); expected exit 0 and "
            f"{expected_ok!r}:\n{output.rstrip()}"
        )
    print(f"OK: review_validate accepts the valid findings file ({expected_ok})")

    problems: list[str] = []
    for label, mutate in REJECT_VECTORS:
        document = copy.deepcopy(VALID_FINDINGS)
        try:
            expected = mutate(document, text_en)
        except KeyError as exc:
            problems.append(
                f"{label}: the vector could not be built — {exc} is no longer in "
                "the accept vector, so this rule is not being tested"
            )
            continue
        path = scratch / "reject.json"
        path.write_text(json.dumps(document, ensure_ascii=False, indent=2),
                        encoding="utf-8")
        code, output = run_quiet("review_validate.py", "--project", project,
                                 "--input", str(path))
        if code != 1:
            problems.append(f"{label}: exit {code}, expected 1")
        elif expected not in output:
            problems.append(
                f"{label}: exit 1 but the reason did not fire — expected "
                f"{expected!r} in:\n      {output.rstrip()}"
            )
        else:
            print(f"  reject {label:<32} exit 1, {expected!r} fires")

    for label, payload, expected in RAW_REJECT_VECTORS:
        path = scratch / "raw.json"
        if payload is None:
            path = scratch / "does-not-exist.json"
            if path.exists():
                path.unlink()
        else:
            path.write_text(payload, encoding="utf-8")
        code, output = run_quiet("review_validate.py", "--project", project,
                                 "--input", str(path))
        if code != 1:
            problems.append(f"{label}: exit {code}, expected 1")
        elif expected not in output:
            problems.append(
                f"{label}: exit 1 but the reason did not fire — expected "
                f"{expected!r} in:\n      {output.rstrip()}"
            )
        else:
            print(f"  reject {label:<32} exit 1, {expected!r} fires")

    if problems:
        raise Failure("review_validate vectors:\n  " + "\n  ".join(problems))
    total = len(REJECT_VECTORS) + len(RAW_REJECT_VECTORS)
    print(f"OK: review_validate — 1 accept vector, {total} reject vectors, "
          "every rejection reason fires")


# --------------------------------------------------------------------------
# review_compare


def build_canned_project(source: str, root: Path, flags: dict) -> None:
    """A project carrying only what review_compare.py reads: a copy of the
    fixture's packet, canned findings and canned flags. Canned so the buckets are
    known exactly, and outside projects/ so no real flag file is touched."""
    (root / "state").mkdir(parents=True, exist_ok=True)
    (root / "review" / "packet").mkdir(parents=True, exist_ok=True)
    shutil.copy2(packet_dir(source) / "pairs.json",
                 root / "review" / "packet" / "pairs.json")
    (root / "state" / "flags.json").write_text(
        json.dumps(flags, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    (root / "review" / "findings.json").write_text(
        json.dumps(CANNED_FINDINGS, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8")


def assert_compare(source: str, scratch: Path) -> None:
    root = scratch / "canned"
    build_canned_project(source, root, CANNED_FLAGS)
    state_dir = root / "state"
    before = state_digest(state_dir)
    run_capture("review_compare.py", "--project", str(root))
    after = state_digest(state_dir)

    problems: list[str] = []
    if after != before:
        problems.append(
            f"review_compare.py wrote to {state_dir} ({before} -> {after}); it is "
            "read-only with respect to state/"
        )
    output_path = root / "review" / "comparison_candidates.json"
    if not output_path.is_file():
        raise Failure(f"review_compare.py wrote no {output_path}")
    doc = json.loads(output_path.read_text(encoding="utf-8"))

    if doc["meta"]["project"] != root.name:
        problems.append(f"meta.project is {doc['meta']['project']!r}, "
                        f"want {root.name!r}")
    if doc["meta"]["counts"] != EXPECTED_COMPARE_COUNTS:
        problems.append(f"meta.counts is {doc['meta']['counts']}, "
                        f"want {EXPECTED_COMPARE_COUNTS}")
    candidates = tuple(
        (c["segment_id"], c["finding_id"], c["flag_key"]) for c in doc["candidates"]
    )
    if candidates != EXPECTED_CANDIDATES:
        problems.append(f"candidates are {candidates}, want {EXPECTED_CANDIDATES}")
    reviewer_only = tuple(
        (r["segment_id"], r["finding_id"]) for r in doc["reviewer_only"]
    )
    if reviewer_only != EXPECTED_REVIEWER_ONLY:
        problems.append(f"reviewer_only is {reviewer_only}, "
                        f"want {EXPECTED_REVIEWER_ONLY}")
    pipeline_only = tuple(
        (p["segment_id"], p["flag_key"]) for p in doc["pipeline_only"]
    )
    if pipeline_only != EXPECTED_PIPELINE_ONLY:
        problems.append(f"pipeline_only is {pipeline_only}, "
                        f"want {EXPECTED_PIPELINE_ONLY}")
    # Verbatim copies, never summaries: the file has to be readable on its own.
    flag_issues = {flag["key"]: flag["issue"] for flag in CANNED_FLAGS["flags"]}
    finding_issues = {f["id"]: f["issue"] for f in CANNED_FINDINGS["findings"]}
    for candidate in doc["candidates"]:
        if candidate["flag_issue"] != flag_issues[candidate["flag_key"]]:
            problems.append(f"{candidate['flag_key']} flag_issue is not verbatim")
        if candidate["finding_issue"] != finding_issues[candidate["finding_id"]]:
            problems.append(f"{candidate['finding_id']} finding_issue is not verbatim")
    if problems:
        raise Failure("review_compare on canned inputs:\n  " + "\n  ".join(problems))
    print("OK: review_compare — 5 candidates (the full cross product on D-0002 "
          "plus the document-wide pair), 2 reviewer_only, 2 pipeline_only, "
          "state/ unchanged")


def assert_compare_no_flags(source: str, scratch: Path) -> None:
    """A project with no flags is a normal input, not an error."""
    root = scratch / "canned-no-flags"
    build_canned_project(source, root, {"flags": []})
    run_capture("review_compare.py", "--project", str(root))
    doc = json.loads(
        (root / "review" / "comparison_candidates.json").read_text(encoding="utf-8")
    )
    problems: list[str] = []
    if doc["candidates"] or doc["pipeline_only"]:
        problems.append("candidates and pipeline_only must both be empty")
    order = tuple(r["finding_id"] for r in doc["reviewer_only"])
    if order != EXPECTED_NO_FLAGS_ORDER:
        problems.append(f"reviewer_only order is {order}, "
                        f"want {EXPECTED_NO_FLAGS_ORDER}")
    if problems:
        raise Failure("review_compare with no flags:\n  " + "\n  ".join(problems))
    print("OK: review_compare — with no flags every finding is reviewer_only, "
          "in packet order with DOC last")


# --------------------------------------------------------------------------
# review_sources (offline mode only)


def _manifest_entry(document: dict, ident: str) -> dict:
    for entry in document["sources"]:
        if entry.get("id") == ident:
            return entry
    raise Failure(
        f"{MANIFEST} has no source {ident!r}: a reject vector targets it, so the "
        "manifest check is no longer being tested"
    )


def assert_manifest(scratch: Path) -> None:
    output = run_capture("review_sources.py", "verify-manifest")
    document = json.loads(MANIFEST.read_text(encoding="utf-8"))
    expected_ok = f"OK: {len(document['sources'])} source(s) consistent"
    if expected_ok not in output:
        raise Failure(f"expected {expected_ok!r} from verify-manifest:\n{output}")
    print(f"OK: review_sources — manifest and style-guide.md Sources agree "
          f"({len(document['sources'])} sources)")

    problems: list[str] = []
    for label, mutate in MANIFEST_REJECT_VECTORS:
        mutated = copy.deepcopy(document)
        expected = mutate(mutated)
        path = scratch / "manifest.json"
        path.write_text(json.dumps(mutated, ensure_ascii=False, indent=2),
                        encoding="utf-8")
        code, output = run_quiet("review_sources.py", "verify-manifest",
                                 "--manifest", str(path))
        if code != 1:
            problems.append(f"{label}: exit {code}, expected 1")
        elif expected not in output:
            problems.append(
                f"{label}: exit 1 but the reason did not fire — expected "
                f"{expected!r} in:\n      {output.rstrip()}"
            )
        else:
            print(f"  reject {label:<34} exit 1, reason fires")
    if problems:
        raise Failure("review_sources reject vectors:\n  " + "\n  ".join(problems))
    print(f"OK: review_sources — {len(MANIFEST_REJECT_VECTORS)} inconsistencies "
          "between manifest and style guide are all detected")


# --------------------------------------------------------------------------
# seeded review fixture


def assert_seeded_defects(project: str) -> None:
    """All eight judgement defects present in the English the reviewer reads.

    Asserted on the packet rather than on state/, because the packet is what a
    blind pass sees: a seed that landed in state but never reached the packet
    would be unmeasurable.
    """
    seeds = json.loads(SEED_TABLE.read_text(encoding="utf-8"))["seeds"]
    text_en = packet_text_en(project)
    text = packet_text(project)

    problems: list[str] = []
    modes = [seed["mode"] for seed in seeds]
    for mode in EXPECTED_SEED_MODES:
        if modes.count(mode) != 1:
            problems.append(f"mode {mode!r} is seeded {modes.count(mode)} time(s), "
                            "want exactly 1")
    for mode in sorted(set(modes) - set(EXPECTED_SEED_MODES)):
        problems.append(f"mode {mode!r} is not one of the eight expected modes")

    for seed in seeds:
        label = f"seed {seed['seed']} ({seed['mode']}) {seed['segment_id']}"
        delivered = text_en.get(seed["segment_id"])
        if delivered is None:
            problems.append(f"{label}: the segment is not in the packet at all")
            continue
        for change in seed["changes"]:
            if change["after"] not in delivered:
                problems.append(
                    f"{label}: the seeded English {change['after'][:60]!r} is not in "
                    "the packet entry"
                )
            if text.count(change["before"]) != 0:
                problems.append(
                    f"{label}: the clean English {change['before'][:60]!r} is still "
                    "in the packet"
                )
    if problems:
        raise Failure(f"seeded defects in {project}:\n  " + "\n  ".join(problems))
    print(f"OK: {project} — all {len(EXPECTED_SEED_MODES)} seeded judgement "
          "defects reach the packet, and no clean fragment survives")


# --------------------------------------------------------------------------


def main() -> int:
    ap = argparse.ArgumentParser(description="Offline smoke test of the pipeline.")
    ap.add_argument(
        "--keep-going", action="store_true",
        help="report every failed assertion instead of stopping at the first",
    )
    args = ap.parse_args()

    if not SCRIPTS.is_dir():
        print(f"ERROR: {SCRIPTS} not found — run from the repo root", file=sys.stderr)
        return 2

    failures: list[str] = []

    def step(fn, *fn_args) -> None:
        try:
            fn(*fn_args)
        except Failure as exc:
            if not args.keep_going:
                raise
            failures.append(str(exc))

    try:
        print("=" * 70)
        print("CLEAN FIXTURE")
        print("=" * 70)
        run("dev/make_fixture.py")
        run("ingest.py", "--project", FIXTURE)
        run("dev/make_fixture_en.py", "--project", FIXTURE)
        run("validate_state.py", "--project", FIXTURE)
        run("glossary.py", "validate", "--project", FIXTURE)
        run("checks.py", "run", "--project", FIXTURE)
        step(assert_statuses, FIXTURE, EXPECTED_CLEAN, "clean fixture")
        # A CONVENTION flag ON A SEGMENT, filed through the real writer. Without
        # it assert_conventions_stay_internal cannot fail: a document-wide flag
        # (segment_id null) never lands on a review-table row, so excluding
        # CONVENTION or not would look identical.
        run("flags.py", "add", "--project", FIXTURE, "--class", "CONVENTION",
            "--segment", "C-01", "--issue", FIXTURE_CONVENTION_ISSUE)
        run("assemble.py", "run", "--project", FIXTURE)
        step(assert_deliverables)
        step(assert_conventions_stay_internal)

        print()
        print("=" * 70)
        print("SEEDED-DEFECT FIXTURE")
        print("=" * 70)
        run("dev/make_fixture_bad.py")
        run("checks.py", "run", "--project", FIXTURE_BAD)
        step(assert_statuses, FIXTURE_BAD, EXPECTED_BAD, "seeded-defect fixture")

        def assert_claim5_named() -> None:
            found = details(FIXTURE_BAD, "antecedent_basis")
            if not any("claim 5" in d for d in found):
                raise Failure(
                    "antecedent_basis warned but did not name claim 5: " + str(found)
                )
            print("OK: antecedent_basis names claim 5")

        step(assert_claim5_named)

        with tempfile.TemporaryDirectory(prefix="patent-smoke-") as tmp:
            scratch = Path(tmp)

            print()
            print("=" * 70)
            print("REVIEW MODULE")
            print("=" * 70)
            step(assert_packet_build, FIXTURE, "clean fixture packet")
            step(assert_packet_shape, FIXTURE, "clean fixture packet")
            step(assert_packet_blind, FIXTURE, "clean fixture packet")
            step(assert_packet_deterministic, FIXTURE, "clean fixture packet")
            step(assert_packet_refuses_stray_file, FIXTURE, "clean fixture packet")
            step(assert_validator, FIXTURE, scratch / "validate")
            step(assert_compare, FIXTURE, scratch / "compare")
            step(assert_compare_no_flags, FIXTURE, scratch / "compare")
            step(assert_manifest, scratch)

            print()
            print("=" * 70)
            print("SEEDED REVIEW FIXTURE")
            print("=" * 70)
            run("dev/seed_review_fixture.py")
            step(assert_packet_build, FIXTURE_REVIEW, "review fixture packet")
            step(assert_packet_shape, FIXTURE_REVIEW, "review fixture packet")
            step(assert_packet_blind, FIXTURE_REVIEW, "review fixture packet")
            step(assert_seeded_defects, FIXTURE_REVIEW)
    except Failure as exc:
        print(f"\nSMOKE FAILED: {exc}", file=sys.stderr)
        return 1

    if failures:
        print("\nSMOKE FAILED:", file=sys.stderr)
        for line in failures:
            print(f"  - {line}", file=sys.stderr)
        return 1

    print("\n" + "=" * 70)
    print("SMOKE OK — the pipeline runs, the checks catch every seeded defect, and")
    print("the review packet stays blind while carrying all eight review defects.")
    print("=" * 70)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
