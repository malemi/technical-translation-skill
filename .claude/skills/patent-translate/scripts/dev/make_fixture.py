#!/usr/bin/env python3
"""Build projects/_fixture/source.docx from the literal fixture text.

The Italian text below is copied EXACTLY from projects/_fixture/fixture_spec.md
(decimal commas, micro sign, numeral suffixes are test data — do not touch).
One paragraph per line; an empty paragraph is inserted between consecutive
text lines so that ingest's empty-paragraph handling is exercised. Run from
the repo root. Output content is deterministic (fixed core properties; zip
member timestamps are whatever zipfile stamps and are not content).
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

from docx import Document

OUT_PATH = Path("projects/_fixture/source.docx")

TEXT_LINES = [
    'TITOLO',
    "Apparecchiatura per l'estrazione a freddo di caffè e relativo metodo di estrazione",
    'RIASSUNTO',
    "Apparecchiatura (1) per l'estrazione a freddo di caffè comprendente una vasca di estrazione (2) atta a contenere acqua e caffè macinato, un gruppo filtrante (3) disposto all'interno della vasca di estrazione (2), e una pompa di ricircolo (4) configurata per far circolare l'acqua attraverso il gruppo filtrante (3). L'apparecchiatura comprende inoltre un'unità di controllo (5) che regola la temperatura dell'acqua tra 4 °C e 8 °C e la pressione di esercizio fino a 12,5 bar. È inoltre descritto un metodo di estrazione a freddo impiegante detta apparecchiatura.",
    'DESCRIZIONE',
    'CAMPO TECNICO',
    "La presente invenzione riguarda un'apparecchiatura per l'estrazione a freddo di caffè, nonché un metodo di estrazione impiegante tale apparecchiatura.",
    'STATO DELLA TECNICA',
    'Sono noti sistemi per la preparazione di caffè mediante estrazione a freddo, nei quali il caffè macinato è posto a contatto con acqua a temperatura ambiente per tempi di estrazione tipicamente compresi tra 12 e 24 ore. Tali sistemi presentano tuttavia consumi energetici elevati e una resa aromatica limitata. Essa risulta inoltre difficilmente controllabile.',
    'BREVE DESCRIZIONE DELLE FIGURE',
    "La figura 1 mostra una vista schematica dell'apparecchiatura secondo l'invenzione.",
    'La figura 2 mostra una sezione del gruppo filtrante.',
    'DESCRIZIONE DETTAGLIATA',
    "Con riferimento alla figura 1, l'apparecchiatura (1) comprende una vasca di estrazione (2) realizzata in acciaio inossidabile, avente una capacità compresa tra 5 e 50 litri.",
    'Il gruppo filtrante (3) comprende una rete metallica con maglie da 100 µm e un supporto (3a) removibile.',
    'La pompa di ricircolo (4) è configurata per operare a una pressione compresa tra 0,5 bar e 12,5 bar, preferibilmente pari a 2,5 bar.',
    "L'unità di controllo (5) comprende un sensore di temperatura (6) e un temporizzatore (7). Il serbatoio (2) è mantenuto a una temperatura compresa tra 4 °C e 8 °C per un tempo di almeno 20 minuti.",
    "In una forma di realizzazione preferita, la percentuale di caffè macinato rispetto all'acqua è compresa tra il 5% e il 12,5% in peso.",
    "Naturalmente, senza pregiudizio per il principio dell'invenzione, i dettagli di realizzazione potranno variare rispetto a quanto descritto.",
    'RIVENDICAZIONI',
    "1. Apparecchiatura (1) per l'estrazione a freddo di caffè, comprendente:",
    'una vasca di estrazione (2) atta a contenere acqua e caffè macinato;',
    "un gruppo filtrante (3) disposto all'interno della vasca di estrazione (2);",
    "una pompa di ricircolo (4) configurata per far circolare l'acqua attraverso il gruppo filtrante (3);",
    "caratterizzata dal fatto che comprende inoltre un'unità di controllo (5) configurata per mantenere la temperatura dell'acqua tra 4 °C e 8 °C.",
    '2. Apparecchiatura secondo la rivendicazione 1, in cui la vasca di estrazione (2) è realizzata in acciaio inossidabile e presenta una capacità compresa tra 5 e 50 litri.',
    '3. Apparecchiatura secondo la rivendicazione 2, in cui il gruppo filtrante (3) comprende una rete metallica con maglie da 100 µm.',
    '4. Apparecchiatura secondo una qualsiasi delle rivendicazioni da 1 a 3, in cui la pompa di ricircolo (4) è costituita da una pompa centrifuga a velocità variabile.',
    "5. Apparecchiatura secondo la rivendicazione precedente, comprendente inoltre un sensore di pressione (8) collegato all'unità di controllo (5).",
    "6. Metodo di estrazione a freddo di caffè mediante un'apparecchiatura secondo una qualsiasi delle rivendicazioni precedenti, comprendente le fasi di:",
    'introdurre acqua e caffè macinato nella vasca di estrazione (2);',
    "far circolare l'acqua mediante la pompa di ricircolo (4) per un tempo di almeno 20 minuti;",
    'mantenere la temperatura tra 4 °C e 8 °C.',
]


def build_paragraphs() -> list[str]:
    paragraphs: list[str] = []
    for i, line in enumerate(TEXT_LINES):
        if i:
            paragraphs.append("")
        paragraphs.append(line)
    return paragraphs


def add_patent_numbering(doc):
    """Number the description paragraphs the way a real application does.

    Italian applications write [0001], [0002], … with Word's automatic list
    numbering, so the digits are in the numbering definition and NOT in the
    paragraph text. Reproducing that here is the point: a fixture whose numbers
    are typed as text would exercise nothing, and this is precisely the path on
    which a whole document's paragraph numbers once went missing unnoticed.

    Returns the labels the numbering renders, in document order.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn

    numbering = doc.part.numbering_part.element
    abstract_id = str(
        max((int(a.get(qn("w:abstractNumId"))) for a in numbering.findall(qn("w:abstractNum"))),
            default=-1) + 1
    )
    num_id = str(
        max((int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))), default=0) + 1
    )
    numbering.append(parse_xml(
        f'<w:abstractNum {nsdecls("w")} w:abstractNumId="{abstract_id}">'
        f'<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimalZero"/>'
        f'<w:lvlText w:val="[00%1]"/><w:lvlJc w:val="left"/></w:lvl></w:abstractNum>'
    ))
    numbering.append(parse_xml(
        f'<w:num {nsdecls("w")} w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abstract_id}"/></w:num>'
    ))

    numbered = set(description_body_lines())
    labels = []
    for para in doc.paragraphs:
        if para.text.strip() not in numbered:
            continue
        para._p.get_or_add_pPr().append(parse_xml(
            f'<w:numPr {nsdecls("w")}><w:ilvl w:val="0"/>'
            f'<w:numId w:val="{num_id}"/></w:numPr>'
        ))
        labels.append(f"[00{len(labels) + 1:02d}]")
    return labels


def description_body_lines() -> list[str]:
    """The description's prose paragraphs — everything between the DESCRIZIONE
    and RIVENDICAZIONI markers that is not itself an all-caps sub-heading."""
    start = TEXT_LINES.index("DESCRIZIONE") + 1
    end = TEXT_LINES.index("RIVENDICAZIONI")
    return [
        line for line in TEXT_LINES[start:end]
        if not (line == line.upper() and not line.endswith("."))
    ]


def main() -> int:
    paragraphs = build_paragraphs()
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    labels = add_patent_numbering(doc)

    props = doc.core_properties
    fixed = datetime(2026, 1, 1, tzinfo=timezone.utc)
    props.created = fixed
    props.modified = fixed
    props.author = "make_fixture"
    props.last_modified_by = "make_fixture"
    props.title = "patent-translate fixture"
    props.revision = 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))

    # self-check: reopen and compare paragraph texts with the intended list
    reopened = [p.text for p in Document(str(OUT_PATH)).paragraphs]
    if reopened != paragraphs:
        raise SystemExit(
            f"ERROR: reopened docx does not match intended paragraphs "
            f"({len(reopened)} vs {len(paragraphs)})"
        )

    # self-check: the numbering we wrote must render to the labels we intended,
    # read back through the same resolver ingest.py uses.
    import sys

    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    import common

    rendered, warnings = common.iter_paragraph_numbers(OUT_PATH)
    if warnings:
        raise SystemExit("ERROR: numbering warnings on the fixture: " + "; ".join(warnings))
    if list(rendered.values()) != labels:
        raise SystemExit(
            f"ERROR: rendered paragraph numbers {list(rendered.values())} "
            f"!= intended {labels}"
        )

    non_empty = sum(1 for p in paragraphs if p)
    print(f"wrote {OUT_PATH}: {len(paragraphs)} paragraphs ({non_empty} non-empty), "
          f"{len(labels)} numbered {labels[0]}–{labels[-1]}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
