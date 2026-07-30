#!/usr/bin/env python3
"""Assemble the English deliverables of a patent-translate project.

Usage: assemble.py run --project projects/<slug>

Reads the finished state (segments, translations, model claims, flags, checks)
and writes, into <project>/out/:
  filing_en.docx      the clean filing (title, description, claims, abstract)
  side_by_side.docx   landscape bilingual review table + audit appendices
  ESCALATIONS.md      open non-convention flags, grouped by class
  audit_numbers.csv   the numbers-and-units audit table (checks data)
  terminology.csv     a copy of the project's glossary
  notes-for-human-reviewer.md   copied from the project when it exists: how a
                      handful of expressions were rendered, and nothing else

This script never runs the checks: checks.py must have produced checks.json
first. It fails loud on any missing required input.
"""

from __future__ import annotations

import argparse
import csv
import re
import shutil
from pathlib import Path

import common

# Appendix section headings in the side-by-side, in render order. Kept as a
# constant because the acceptance test asserts each one is present.
APPENDIX_HEADINGS = (
    "Numbers and units audit",
    "Reference numeral map",
    "Claim dependency map",
    "Open escalations",
)

# What a human editor is told about drafting conventions: notes-for-human-reviewer.md
# and nothing else. CONVENTION flags are an internal record and are kept out of the
# side-by-side, out of ESCALATIONS.md and out of the appendices — one per
# occurrence of a transitional phrase put 18 rows of noise in cafe124's review
# table, which is 18 more places for a reader to lose the findings that matter.
NOTES_FOR_REVIEWER = "notes-for-human-reviewer.md"

# Classes that escalate to ESCALATIONS.md (everything open except CONVENTION).
ESCALATION_CLASSES = ("AMBIGUITY", "TERM", "CLAIM-DEFECT", "MECH")

_CLAIM_NUMBER_RE = re.compile(r"^\s*\d+\.\s")


# --------------------------------------------------------------------------
# input loading and gap checks


def _en_of_pair(pair: dict) -> str | None:
    """Chosen English of a translation pair: final if set, else DeepL."""
    final = pair.get("text_en_final")
    return final if final is not None else pair.get("text_en_deepl")


def load_state(paths: dict) -> dict:
    """Load every required state file, failing loud on anything missing."""
    for label, path in (
        ("segments.json", paths["segments"]),
        ("translations.json", paths["translations"]),
        ("claims_en.json", paths["claims_en"]),
        ("flags.json", paths["flags"]),
        ("checks.json (run checks.py first)", paths["checks"]),
    ):
        if not path.is_file():
            common.die(f"required state file missing: {label} -> {path}")
    if not paths["terminology"].is_file():
        common.die(f"required terminology.csv missing: {paths['terminology']}")
    return {
        "segments": common.read_json(paths["segments"]),
        "translations": common.read_json(paths["translations"]),
        "claims_en": common.read_json(paths["claims_en"]),
        "flags": common.read_json(paths["flags"]),
        "checks": common.read_json(paths["checks"]),
    }


def assert_no_gaps(translations: dict) -> None:
    """Every pair must have usable English and the title must be translated."""
    gaps: list[str] = []
    if translations.get("title", {}).get("text_en") is None:
        gaps.append("title.text_en is not set")
    for pair in translations["pairs"]:
        en = _en_of_pair(pair)
        if en is None or not str(en).strip():
            gaps.append(f"pair {pair['id']} has no English (final or DeepL)")
    if gaps:
        common.die(
            "cannot assemble; translations are incomplete:\n  - "
            + "\n  - ".join(gaps)
        )


# --------------------------------------------------------------------------
# per-segment English lookup


def build_english_index(state: dict):
    """Return (english_for, pairs_by_id, claims_by_id, flags_by_segment).

    english_for(segment) yields the chosen English for any segment (title,
    heading, abstract, description via pairs; claims via claims_en); returns
    "" for segments never translated (e.g. the TITOLO/RIVENDICAZIONI markers).
    """
    title_en = state["translations"]["title"]["text_en"]
    pairs_by_id = {p["id"]: p for p in state["translations"]["pairs"]}
    claims_by_id = {c["id"]: c for c in state["claims_en"]["claims"]}

    flags_by_segment: dict[str, list] = {}
    for flag in state["flags"]["flags"]:
        seg_id = flag.get("segment_id")
        if seg_id is not None:
            flags_by_segment.setdefault(seg_id, []).append(flag)

    def english_for(segment: dict) -> str:
        seg_id = segment["id"]
        if segment["kind"] == "title":
            return title_en or ""
        if segment["kind"] == "claim":
            claim = claims_by_id.get(seg_id)
            if claim is None:
                common.die(f"claim segment {seg_id} has no claims_en entry")
            return claim["text_en"]
        pair = pairs_by_id.get(seg_id)
        if pair is not None:
            return _en_of_pair(pair) or ""
        return ""  # section markers (title/claims headings) are never sent

    return english_for, pairs_by_id, claims_by_id, flags_by_segment


# --------------------------------------------------------------------------
# filing_en.docx


def build_filing(state: dict, english_for, out_path) -> int:
    """Clean English filing. Abstract is rendered ONCE, at the very end.

    Returns the number of claims rendered (for the reopen assertion).
    """
    from docx import Document

    doc = Document()
    title_en = state["translations"]["title"]["text_en"]
    doc.add_paragraph(title_en, style="Heading 1")

    # Body: description section only. Title- and abstract-section headings are
    # deliberately skipped here (the title is the Heading 1 above and the
    # abstract is rendered once at the end); rendering them would leave orphan
    # or duplicate headings. See report note.
    for segment in state["segments"]["segments"]:
        if segment["section"] != "description":
            continue
        text = english_for(segment)
        if segment["kind"] == "heading":
            doc.add_paragraph(text, style="Heading 2")
        else:
            doc.add_paragraph(text, style="Normal")

    # Claims.
    doc.add_paragraph("CLAIMS", style="Heading 2")
    claims = sorted(state["claims_en"]["claims"], key=lambda c: c["number"])
    for claim in claims:
        parts = claim["parts_en"]
        if not parts:
            common.die(f"claim {claim['id']} has empty parts_en")
        doc.add_paragraph(f"{claim['number']}. {parts[0]}", style="Normal")
        for part in parts[1:]:
            doc.add_paragraph(part, style="Normal")

    # Abstract, once, at the end.
    doc.add_paragraph("ABSTRACT", style="Heading 2")
    for segment in state["segments"]["segments"]:
        if segment["kind"] == "abstract":
            doc.add_paragraph(english_for(segment), style="Normal")

    doc.save(str(out_path))
    return len(claims)


# --------------------------------------------------------------------------
# side_by_side.docx


def _set_landscape(doc) -> None:
    from docx.enum.section import WD_ORIENT

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    width, height = section.page_width, section.page_height
    if width < height:  # default template is portrait; swap to widen the page
        section.page_width, section.page_height = height, width


def _flag_lines(flags: list) -> list[str]:
    """The review table's Flags column: everything except CONVENTION.

    A CONVENTION flag is a rule we applied and disclosed, not a doubt. The
    disclosure belongs in notes-for-human-reviewer.md, once per rule; repeating
    it on every row it touches buries the flags that are actually questions.
    """
    return [
        f"⚑ {f['class']} — {f['issue']} [{f['status']}]"
        for f in flags if f["class"] != "CONVENTION"
    ]


def _fill_cell(cell, lines: list[str]) -> None:
    cell.text = lines[0] if lines else ""
    for extra in lines[1:]:
        cell.add_paragraph(extra)


def _add_table(doc, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, headers):
        cell.text = head
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = "" if value is None else str(value)
    return table


def _fmt_number_units(items) -> str:
    """Render a list of (number, unit) pairs as "12.5 bar; 4 °C"."""
    parts = []
    for item in items or []:
        number = item[0]
        unit = item[1] if len(item) > 1 else None
        parts.append(f"{number} {unit}" if unit else f"{number}")
    return "; ".join(parts) if parts else "—"


def _fmt_list(values) -> str:
    return ", ".join(str(v) for v in values) if values else "—"


def _check_data(checks: dict, name: str):
    """Return a check's table rows (its data['rows']) or None.

    checks.py stores per-check tables as {"rows": [...]}; every caller here
    iterates the rows, so unwrap the envelope once."""
    for result in checks["results"]:
        if result["check"] == name:
            data = result.get("data")
            if isinstance(data, dict):
                return data.get("rows")
            return data
    return None


def build_side_by_side(state: dict, english_for, flags_by_segment, out_path):
    from docx import Document

    doc = Document()
    _set_landscape(doc)
    doc.add_paragraph("Bilingual review", style="Heading 1")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    for cell, head in zip(table.rows[0].cells, ("ID", "Italiano", "English", "Flags")):
        cell.text = head
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for segment in state["segments"]["segments"]:
        cells = table.add_row().cells
        cells[0].text = segment["id"]
        cells[1].text = segment["text_it"]
        cells[2].text = english_for(segment)
        _fill_cell(cells[3], _flag_lines(flags_by_segment.get(segment["id"], [])))

    _append_appendices(doc, state, flags_by_segment)
    doc.save(str(out_path))


def _append_appendices(doc, state: dict, flags_by_segment) -> None:
    checks = state["checks"]

    # 1) Numbers and units audit (checks data 4: numbers_units).
    doc.add_paragraph(APPENDIX_HEADINGS[0], style="Heading 2")
    numbers = _check_data(checks, "numbers_units")
    if numbers:
        rows = [
            [
                row.get("segment", ""),
                _fmt_number_units(row.get("it")),
                _fmt_number_units(row.get("en")),
                "yes" if row.get("ok") else "no",
            ]
            for row in numbers
        ]
        _add_table(doc, ["Segment", "Italiano", "English", "Match"], rows)
    else:
        doc.add_paragraph("(no numbers_units data)", style="Normal")

    # 2) Reference numeral map (checks data 3: numeral_term_consistency).
    doc.add_paragraph(APPENDIX_HEADINGS[1], style="Heading 2")
    numerals = _check_data(checks, "numeral_term_consistency")
    if numerals:
        rows = [
            [
                row.get("numeral", ""),
                _fmt_list(row.get("it_terms")),
                _fmt_list(row.get("en_terms")),
            ]
            for row in numerals
        ]
        _add_table(doc, ["Numeral", "Italian terms", "English terms"], rows)
    else:
        doc.add_paragraph("(no numeral_term_consistency data)", style="Normal")

    # 3) Claim dependency map (checks data 5: claims_graph).
    doc.add_paragraph(APPENDIX_HEADINGS[2], style="Heading 2")
    graph = _check_data(checks, "claims_graph")
    if graph:
        rows = [
            [
                str(row.get("number", "")),
                _fmt_list(row.get("it_deps")),
                _fmt_list(row.get("en_deps")),
                "yes" if row.get("multiple") else "no",
            ]
            for row in graph
        ]
        _add_table(doc, ["Claim", "IT deps", "EN deps", "Multiple"], rows)
    else:
        doc.add_paragraph("(no claims_graph data)", style="Normal")

    # 4) Open escalations (open flags, class != CONVENTION).
    doc.add_paragraph(APPENDIX_HEADINGS[3], style="Heading 2")
    open_esc = [
        f
        for flags in flags_by_segment.values()
        for f in flags
        if f["status"] == "open" and f["class"] != "CONVENTION"
    ]
    if open_esc:
        for flag in open_esc:
            doc.add_paragraph(
                f"{flag.get('segment_id', '?')} — {flag['class']}: {flag['issue']}",
                style="Normal",
            )
            if flag.get("options"):
                doc.add_paragraph("Options: " + "; ".join(flag["options"]), style="Normal")
    else:
        doc.add_paragraph("None.", style="Normal")

    # There is deliberately no "Applied conventions" appendix. See
    # NOTES_FOR_REVIEWER at the top of this module.


# --------------------------------------------------------------------------
# ESCALATIONS.md and audit_numbers.csv


def build_escalations_md(state: dict, out_path) -> None:
    by_class: dict[str, list] = {cls: [] for cls in ESCALATION_CLASSES}
    for flag in state["flags"]["flags"]:
        if flag["status"] == "open" and flag["class"] in by_class:
            by_class[flag["class"]].append(flag)

    lines = ["# Escalations", ""]
    any_flag = False
    for cls in ESCALATION_CLASSES:
        flags = by_class[cls]
        if not flags:
            continue
        any_flag = True
        lines.append(f"## {cls}")
        lines.append("")
        for flag in flags:
            lines.append(f"### {flag.get('segment_id', '(no segment)')}")
            lines.append("")
            lines.append(f"- **IT:** {flag.get('text_it') or '(none)'}")
            lines.append(f"- **EN:** {flag.get('text_en') or '(none)'}")
            lines.append(f"- **Issue:** {flag['issue']}")
            if flag.get("options"):
                lines.append("- **Options:**")
                for option in flag["options"]:
                    lines.append(f"  - {option}")
            lines.append("")
    if not any_flag:
        lines.append("_No open escalations._")
        lines.append("")
    out_path.write_text("\n".join(lines), encoding="utf-8")


def build_audit_csv(state: dict, out_path) -> None:
    numbers = _check_data(state["checks"], "numbers_units") or []
    with open(out_path, "w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(["segment", "italiano", "english", "match"])
        for row in numbers:
            writer.writerow(
                [
                    row.get("segment", ""),
                    _fmt_number_units(row.get("it")),
                    _fmt_number_units(row.get("en")),
                    "yes" if row.get("ok") else "no",
                ]
            )


# --------------------------------------------------------------------------
# reopen-and-assert (contract acceptance, run by assemble itself)


def reopen_and_assert(filing_path, sbs_path, n_claims: int, n_segments: int) -> None:
    from docx import Document

    filing = Document(str(filing_path))
    paras = [(p.style.name, p.text) for p in filing.paragraphs]
    has_claims_heading = any(
        style == "Heading 2" and text.strip() == "CLAIMS" for style, text in paras
    )
    if not has_claims_heading:
        common.die("acceptance failed: filing_en.docx has no CLAIMS heading")

    # Count claim number-prefixed paragraphs between CLAIMS and ABSTRACT.
    in_claims = False
    counted = 0
    for style, text in paras:
        if style == "Heading 2" and text.strip() == "CLAIMS":
            in_claims = True
            continue
        if style == "Heading 2" and text.strip() == "ABSTRACT":
            in_claims = False
            continue
        if in_claims and _CLAIM_NUMBER_RE.match(text):
            counted += 1
    if counted != n_claims:
        common.die(
            f"acceptance failed: filing shows {counted} numbered claims, expected {n_claims}"
        )

    sbs = Document(str(sbs_path))
    if not sbs.tables:
        common.die("acceptance failed: side_by_side.docx has no table")
    row_count = len(sbs.tables[0].rows)
    if row_count != n_segments + 1:
        common.die(
            f"acceptance failed: side-by-side has {row_count} rows, "
            f"expected {n_segments + 1} (segments {n_segments} + header)"
        )
    headings = {p.text.strip() for p in sbs.paragraphs if p.style.name == "Heading 2"}
    missing = [h for h in APPENDIX_HEADINGS if h not in headings]
    if missing:
        common.die("acceptance failed: side-by-side missing appendices: " + ", ".join(missing))


# --------------------------------------------------------------------------
# entry point


def copy_notes_for_reviewer(paths) -> Path | None:
    """Copy <project>/notes-for-human-reviewer.md into out/, or None if absent.

    Authored by the model, not generated: it is the one thing a human editor is
    told about drafting conventions, and it says only how an expression was
    rendered — the Italian, the English, nothing else. No reasons, no doubts, no
    "confirm". A stale copy is worse than none, so an out/ copy left over from a
    run where the project file has since been removed is deleted.
    """
    source = Path(paths["root"]) / NOTES_FOR_REVIEWER
    target = paths["out"] / NOTES_FOR_REVIEWER
    if not source.is_file():
        if target.is_file():
            target.unlink()
        return None
    shutil.copyfile(source, target)
    return target


def run(args) -> int:
    paths = common.project_paths(args.project, need_source=False)
    state = load_state(paths)
    assert_no_gaps(state["translations"])

    english_for, _pairs, _claims, flags_by_segment = build_english_index(state)

    paths["out"].mkdir(parents=True, exist_ok=True)
    filing_path = paths["out"] / "filing_en.docx"
    sbs_path = paths["out"] / "side_by_side.docx"
    escalations_path = paths["out"] / "ESCALATIONS.md"
    audit_path = paths["out"] / "audit_numbers.csv"
    terminology_path = paths["out"] / "terminology.csv"

    n_claims = build_filing(state, english_for, filing_path)
    build_side_by_side(state, english_for, flags_by_segment, sbs_path)
    build_escalations_md(state, escalations_path)
    build_audit_csv(state, audit_path)
    shutil.copyfile(paths["terminology"], terminology_path)
    notes_path = copy_notes_for_reviewer(paths)

    n_segments = len(state["segments"]["segments"])
    reopen_and_assert(filing_path, sbs_path, n_claims, n_segments)

    print("Assembled deliverables:")
    written = [filing_path, sbs_path, escalations_path, audit_path, terminology_path]
    if notes_path is not None:
        written.append(notes_path)
    for path in written:
        print(f"  {path}")
    if notes_path is None:
        print(f"  ({NOTES_FOR_REVIEWER} not authored for this project)")
    print(
        f"Acceptance OK: filing has CLAIMS heading + {n_claims} claims; "
        f"side-by-side rows = {n_segments} segments + 1 header."
    )
    return 0


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)
    run_parser = sub.add_parser("run", help="build the deliverables")
    run_parser.add_argument(
        "--project", required=True, help="project directory, e.g. projects/_fixture"
    )
    args = parser.parse_args()
    if args.command == "run":
        return run(args)
    parser.error(f"unknown command: {args.command}")
    return 2


if __name__ == "__main__":
    raise SystemExit(main())
